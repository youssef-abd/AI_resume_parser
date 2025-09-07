from typing import List, Optional
from uuid import uuid4
from io import BytesIO
import re
import json
import os
import unicodedata
from pathlib import Path
from datetime import datetime

from fastapi import FastAPI, File, Form, HTTPException, UploadFile, Request, Response
from fastapi.responses import PlainTextResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field

app = FastAPI(
    title="AI Resume Parser API",
    version="0.1.0",
    description="API for parsing resumes, extracting skills, and matching with job descriptions.",
)

import pdfplumber
import docx
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import spacy
from spacy.matcher import PhraseMatcher
from sentence_transformers import SentenceTransformer
from app.db import init_db, insert_resume, insert_job, get_job, search_resumes_by_embedding, get_session


class JobIn(BaseModel):
    title: str
    description: str
    required_skills: Optional[List[str]] = Field(default=None, description="Optional canonical skill names")


class UploadResumeResponse(BaseModel):
    resume_id: str
    candidate_name: Optional[str]
    skills: List[str]
    embedding_dim: int
    raw_text_len: int
    cleaned_text_len: int
    note: str


class UploadJobResponse(BaseModel):
    job_id: str
    title: str
    required_skills: List[str]
    embedding_dim: int
    note: str


class JobOut(BaseModel):
    job_id: str
    title: str
    description: str
    required_skills: List[str]


class ResumeOut(BaseModel):
    resume_id: str
    candidate_name: Optional[str]
    cleaned_text: str
    skills: List[str]


class JobListItem(BaseModel):
    job_id: str
    title: str
    created_at: str
    updated_at: str
    required_skills_count: int


class ResumeListItem(BaseModel):
    resume_id: str
    candidate_name: Optional[str]
    skills: List[str]
    created_at: str
    updated_at: str


class MatchResult(BaseModel):
    resume_id: str
    candidate_name: Optional[str]
    cosine: float
    skills_overlap: float
    composite_score: float
    matched_skills: List[str]
    missing_skills: List[str]
    matched_spans: List[dict] = Field(default_factory=list, description="List of {skill, text, start, end} occurrences")
    context_terms: List[str] = Field(default_factory=list, description="Non-skill overlapping terms between job and resume")
    context_job_spans: List[dict] = Field(default_factory=list, description="Spans in job description for context terms")
    context_resume_spans: List[dict] = Field(default_factory=list, description="Spans in resume for context terms")


class MatchResponse(BaseModel):
    job_id: str
    k: int
    results: List[MatchResult]
    note: str


def normalize_text(text: str) -> str:
    """Normalize whitespace and fix common resume formatting artifacts."""
    if not text:
        return ""
    t = unicodedata.normalize("NFKC", text)
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    # De-hyphenate words broken across line breaks, e.g., "micro-\nservices" -> "microservices"
    t = re.sub(r"(\w)-\n(\w)", r"\1\2", t)
    # Collapse spaces/tabs
    t = re.sub(r"[ \t]+", " ", t)
    # Collapse many newlines to max two
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def strip_markdown_and_bullets(text: str) -> str:
    """Remove common markdown syntax, list bullets, and decorative lines while preserving content.
    - Strips leading markdown headers (#, ##) but keeps heading text
    - Removes list bullets (-, *, +, •, –, —) and numbering (1., 1), a), i.) at line starts
    - Unwraps inline formatting (**bold**, *italic*, `code`) and code fences ```...```
    - Converts [label](url) to label; ![alt](url) to alt
    - Removes standalone decorative lines (e.g., -----, ******)
    """
    if not text:
        return ""

    t = text
    # Remove fenced code markers but keep content inside
    t = re.sub(r"```([\s\S]*?)```", r"\1", t)
    # Inline code backticks
    t = re.sub(r"`([^`]*)`", r"\1", t)

    # Images: keep alt text
    t = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", t)
    # Links: keep label
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)

    # Bold/italic markers
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"__([^_]+)__", r"\1", t)
    t = re.sub(r"\*([^*]+)\*", r"\1", t)
    t = re.sub(r"_([^_]+)_", r"\1", t)

    # Heading markers at start of line
    t = re.sub(r"^\s*#{1,6}\s*", "", t, flags=re.MULTILINE)

    # Decorative lines consisting of bullets/asterisks/dashes only
    t = re.sub(r"^\s*[-*_•·]{3,}\s*$", "", t, flags=re.MULTILINE)

    # Bullets at start of line (common unicode bullets and dashes)
    t = re.sub(r"^\s*[\-\*\+\u2022\u2023\u25E6\u2013\u2014\u00B7\u25AA\u25CF]\s+", "", t, flags=re.MULTILINE)

    # Numbered/lettered lists like: 1.  a)  i.  (A)  (iv)
    t = re.sub(r"^\s*\(?([0-9]+|[A-Za-z]|[ivxlcdmIVXLCDM]+)[\).]\s+", "", t, flags=re.MULTILINE)

    return t


def clean_text_pipeline(text: str) -> str:
    """Full cleaning: normalize unicode/newlines, strip markdown/bullets, collapse whitespace."""
    t = normalize_text(text)
    t = strip_markdown_and_bullets(t)
    t = normalize_text(t)
    return t


def extract_text_from_pdf(data: bytes) -> tuple[str, int]:
    """Extract text and page count from a PDF byte buffer."""
    with pdfplumber.open(BytesIO(data)) as pdf:
        texts: List[str] = []
        for page in pdf.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            texts.append(txt)
        raw = "\n".join(texts).strip()
        return raw, len(pdf.pages)


def extract_text_from_docx(data: bytes) -> str:
    """Extract text from a DOCX byte buffer (paragraphs + tables)."""
    document = docx.Document(BytesIO(data))
    parts: List[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join([p for p in parts if p]).strip()


def save_parsed_json(
    resume_id: str,
    candidate_name: Optional[str],
    meta: dict,
    raw_text: str,
    cleaned_text: str,
) -> None:
    outdir = Path("data") / "parsed"
    outdir.mkdir(parents=True, exist_ok=True)
    payload = {
        "resume_id": resume_id,
        "candidate_name": candidate_name,
        "meta": meta,
        "raw_text": raw_text,
        "cleaned_text": cleaned_text,
    }
    with (outdir / f"{resume_id}.json").open("w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


# --- Skill registry loading and extraction ---
import re as _re

SKILLS_REGISTRY_PATH = (Path(__file__).parent / "skills_registry.json").resolve()
CANONICAL_SKILLS: list[str] = []
SKILL_ALIASES: dict[str, str] = {}


def load_skills_registry() -> None:
    global CANONICAL_SKILLS, SKILL_ALIASES
    try:
        with open(SKILLS_REGISTRY_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        skills = data.get("skills", [])
        canon: list[str] = []
        aliases: dict[str, str] = {}
        for item in skills:
            name = (item.get("name") or "").strip().lower()
            if not name:
                continue
            canon.append(name)
            for a in item.get("aliases", []) or []:
                a_norm = (a or "").strip().lower()
                if a_norm:
                    aliases[a_norm] = name
        # Fallback if empty
        if not canon:
            canon = ["python", "sql"]
        CANONICAL_SKILLS = canon
        SKILL_ALIASES = aliases
    except Exception:
        # Fallback minimal
        CANONICAL_SKILLS = ["python", "sql"]
        SKILL_ALIASES = {}


load_skills_registry()

# Precompile regex patterns with word boundaries (case-insensitive)
_SKILL_PATTERNS: dict[str, _re.Pattern] = {}
for skill in CANONICAL_SKILLS:
    pattern = _re.compile(rf"(?<!\w){_re.escape(skill)}(?!\w)", _re.IGNORECASE)
    _SKILL_PATTERNS[skill] = pattern
for alias, canon in SKILL_ALIASES.items():
    pattern = _re.compile(rf"(?<!\w){_re.escape(alias)}(?!\w)", _re.IGNORECASE)
    _SKILL_PATTERNS[alias] = pattern  # map alias -> canon after match


def extract_skills_simple(text: str) -> list[str]:
    """Detect presence of known skills using regex; returns canonical, unique, sorted list."""
    if not text:
        return []
    found: set[str] = set()
    for token, pat in _SKILL_PATTERNS.items():
        if pat.search(text):
            canon = SKILL_ALIASES.get(token, token)
            if canon in CANONICAL_SKILLS:
                found.add(canon)
    return sorted(found)

# --- spaCy PhraseMatcher-based skill extraction ---
try:
    NLP = spacy.load("en_core_web_sm", disable=["ner", "textcat"])
except Exception:  # safe fallback
    NLP = spacy.blank("en")

SKILL_MATCHER = PhraseMatcher(NLP.vocab, attr="LOWER")
# phrase -> canonical mapping
SKILL_CANON_BY_PHRASE: dict[str, str] = {}

def build_skill_matcher() -> None:
    phrases: list[spacy.tokens.Doc] = []
    SKILL_CANON_BY_PHRASE.clear()
    if not CANONICAL_SKILLS:
        return
    for canon in CANONICAL_SKILLS:
        doc = NLP.make_doc(canon)
        phrases.append(doc)
        SKILL_CANON_BY_PHRASE[canon.lower()] = canon
    for alias, canon in SKILL_ALIASES.items():
        doc = NLP.make_doc(alias)
        phrases.append(doc)
        SKILL_CANON_BY_PHRASE[alias.lower()] = canon
    if phrases:
        # Remove existing patterns if any, then add
        try:
            SKILL_MATCHER.remove("SKILLS")
        except Exception:
            pass
        SKILL_MATCHER.add("SKILLS", phrases)

build_skill_matcher()

def extract_context_terms(job_text: str, resume_text: str, exclude_terms: Optional[List[str]] = None, max_terms: int = 20) -> List[str]:
    """Find overlapping context terms (noun chunks + content tokens) between job and resume, excluding provided terms and skills."""
    if not job_text or not resume_text:
        return []
    doc_j = NLP(job_text)
    doc_r = NLP(resume_text)

    def terms_from_doc(doc: spacy.tokens.Doc) -> set[str]:
        terms: set[str] = set()
        # noun chunks (if available)
        for nc in getattr(doc, "noun_chunks", []):
            t = nc.text.strip().lower()
            if 3 <= len(t) <= 60 and not t.isnumeric():
                terms.add(t)
        # content tokens
        for tok in doc:
            if tok.is_stop or tok.is_punct or not tok.text.isalpha():
                continue
            if tok.pos_ in {"NOUN", "PROPN", "ADJ"} and len(tok.text) >= 3:
                terms.add(tok.lemma_.lower())
        return terms

    j_terms = terms_from_doc(doc_j)
    r_terms = terms_from_doc(doc_r)
    overlap = j_terms & r_terms
    exclude = set((exclude_terms or [])) | set(CANONICAL_SKILLS)
    filtered = [t for t in overlap if t not in exclude]
    filtered.sort(key=lambda x: (-len(x), x))  # rough importance: longer terms first
    return filtered[:max_terms]


def find_spans_for_terms(text: str, terms: List[str]) -> List[dict]:
    spans: List[dict] = []
    if not text or not terms:
        return spans
    for t in terms:
        try:
            pat = _re.compile(rf"(?<!\w){_re.escape(t)}(?!\w)", _re.IGNORECASE)
            for m in pat.finditer(text):
                spans.append({"text": m.group(0), "start": int(m.start()), "end": int(m.end())})
        except Exception:
            continue
    return spans


def extract_skills_with_spans(text: str) -> tuple[list[str], list[dict]]:
    if not text:
        return [], []
    doc = NLP(text)
    found: set[str] = set()
    spans: list[dict] = []
    for _, start, end in SKILL_MATCHER(doc):
        span = doc[start:end]
        phrase = span.text.lower()
        canon = SKILL_CANON_BY_PHRASE.get(phrase, phrase)
        if canon in CANONICAL_SKILLS:
            found.add(canon)
            spans.append({
                "skill": canon,
                "text": span.text,
                "start": int(span.start_char),
                "end": int(span.end_char),
            })
    # fallback to regex for anything missed (no spans via regex)
    for s in extract_skills_simple(text):
        found.add(s)
    skills_sorted = sorted(found)
    # keep only spans whose skill is in skills_sorted
    spans = [sp for sp in spans if sp["skill"] in skills_sorted]
    return skills_sorted, spans


def extract_skills(text: str) -> list[str]:
    skills, _ = extract_skills_with_spans(text)
    return skills

# Embedding model (sentence-transformers)
try:
    EMBEDDER = SentenceTransformer("all-MiniLM-L6-v2")
    EMBEDDING_DIM = EMBEDDER.get_sentence_embedding_dimension()
except Exception:
    EMBEDDER = None
    EMBEDDING_DIM = 384

def embed_text(text: str) -> list[float]:
    if not text:
        return [0.0] * EMBEDDING_DIM
    if EMBEDDER is None:
        # zero vector fallback
        return [0.0] * EMBEDDING_DIM
    vec = EMBEDDER.encode([text], normalize_embeddings=True)[0]
    return vec.tolist()

# DB-backed persistence; in-memory stores removed

def normalize_skill_list(skills: list[str] | None) -> list[str]:
    normalized: list[str] = []
    seen: set[str] = set()
    for s in skills or []:
        t = (s or "").strip().lower()
        if not t:
            continue
        canon = SKILL_ALIASES.get(t, t)
        if canon not in seen:
            seen.add(canon)
            normalized.append(canon)
    return normalized


def parse_skill_str_list(value: Optional[str]) -> list[str]:
    if not value:
        return []
    parts = re.split(r'[\n,;]+', value)
    return [p.strip() for p in parts if p.strip()]


def debug_mode() -> bool:
    return (os.getenv("DEBUG") or "").lower() in {"1", "true", "yes", "on"}


def get_max_upload_bytes() -> int:
    """Max upload size for resumes in bytes (default 10 MB). Configurable via MAX_UPLOAD_MB env var."""
    try:
        mb = int(os.getenv("MAX_UPLOAD_MB", "10"))
    except Exception:
        mb = 10
    return max(1, mb) * 1024 * 1024


def jaccard(a: set[str], b: set[str]) -> float:
    if not a and not b:
        return 0.0
    inter = len(a & b)
    union = len(a | b)
    return float(inter) / union if union else 0.0




# CORS configuration for browser clients (Axios)
# FRONTEND_ORIGINS can be set to a comma-separated list of exact origins (e.g., "https://youssef2106-ai-resume-parser.hf.space")
# If FRONTEND_ORIGINS="*" (default), credentials are disabled to satisfy browser rules.
_frontend_origins = os.getenv("FRONTEND_ORIGINS", "*")
if _frontend_origins == "*":
    _allow_origins = ["*"]
    _allow_credentials = False
else:
    _allow_origins = [o.strip() for o in _frontend_origins.split(",") if o.strip()]
    _allow_credentials = True
app.add_middleware(
    CORSMiddleware,
    allow_origins=_allow_origins,
    allow_credentials=_allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)

# Minimal request logger for diagnosing CORS/proxy issues
@app.middleware("http")
async def _log_requests(req: Request, call_next):
    try:
        import time as _t
        t0 = _t.time()
        resp = await call_next(req)
        dt = int((_t.time() - t0) * 1000)
        print(f"[REQ] {req.method} {req.url.path} -> {resp.status_code} {dt}ms origin={req.headers.get('origin','-')}")
        return resp
    except Exception as _e:
        print(f"[REQ] error: {_e}")
        raise

# Initialize database and extensions on startup
@app.on_event("startup")
def _startup() -> None:
    init_db()


@app.get("/healthz")
def healthz():
    return {"status": "ok", "version": app.version}


@app.get("/readyz")
def readyz():
    # Check model loaded
    model_loaded = EMBEDDER is not None and EMBEDDING_DIM > 0
    # Check DB connectivity and basic schema
    db_ok = False
    db_info = {}
    try:
        from sqlalchemy import text as _sql_text
        with get_session() as s:
            ver = s.execute(_sql_text("SELECT version()")).scalar()
            has_vec = s.execute(_sql_text("SELECT EXISTS (SELECT 1 FROM pg_extension WHERE extname='vector')")).scalar()
            resumes_ok = s.execute(_sql_text("SELECT to_regclass('public.resumes')")).scalar() is not None
            jobs_ok = s.execute(_sql_text("SELECT to_regclass('public.jobs')")).scalar() is not None
            db_ok = True
            db_info = {
                "version": ver,
                "has_vector": bool(has_vec),
                "tables": {"resumes": bool(resumes_ok), "jobs": bool(jobs_ok)},
            }
    except Exception as e:
        db_ok = False
        db_info = {"error": str(e)}

    status = "ok" if (model_loaded and db_ok) else "degraded"
    return {"status": status, "model_loaded": model_loaded, "db_ok": db_ok, "db": db_info, "version": app.version}


@app.get("/")
async def root():
    return {"status": "running", "message": "AI Resume Parser API"}


@app.get("/health")
async def health():
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat() + "Z"}


@app.get("/api/health")
async def api_health():
    return {"status": "healthy", "service": "api"}


@app.get("/api/test")
async def api_test():
    return {"status": "ok", "ts": datetime.utcnow().isoformat() + "Z"}

@app.get("/test-deployment")
async def test_deployment():
    return {"message": "New deployment working", "timestamp": "2025-09-07"}


@app.get("/static/js/{filename:path}")
async def serve_javascript(filename: str):
    """
    Serve JavaScript files with correct MIME type to fix browser module loading issues
    """
    # Generate a minimal valid JavaScript module
    js_content = f"""
// Generated JavaScript module for {filename}
// This prevents MIME type errors in HuggingFace Spaces

console.warn('JavaScript module "{filename}" not available in this environment');

// Export empty default to prevent module loading errors
export default {{}};

// If this is the main Streamlit index file, provide basic functionality
if ('{filename}'.includes('index.')) {{
    // Minimal Streamlit compatibility
    window.streamlit = window.streamlit || {{}};
    
    // Suppress common errors
    window.addEventListener('error', function(e) {{
        if (e.message && (
            e.message.includes('Failed to load module') ||
            e.message.includes('MIME type')
        )) {{
            e.preventDefault();
            console.warn('Suppressed module loading error:', e.message);
        }}
    }});
    
    console.info('Streamlit JavaScript compatibility layer loaded');
}}
"""
    
    return Response(
        content=js_content,
        media_type="application/javascript",
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        }
    )

@app.get("/static/css/{filename:path}")
async def serve_css(filename: str):
    """
    Serve CSS files with correct MIME type
    """
    css_content = f"""
/* Generated CSS for {filename} */
/* This prevents MIME type errors in HuggingFace Spaces */

/* Basic Streamlit styling fallback */
body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', sans-serif;
    margin: 0;
    padding: 0;
}}

#root {{
    min-height: 100vh;
}}
"""
    
    return Response(
        content=css_content,
        media_type="text/css",
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )

# Health check endpoint
@app.get("/health")
async def health_check():
    return {"status": "ok", "service": "fastapi"}


@app.post("/upload_resume", response_model=UploadResumeResponse, tags=["resumes"])  # type: ignore[arg-type]
async def upload_resume(
    file: UploadFile = File(..., description="Resume file: PDF or DOCX"),
    candidate_name: Optional[str] = Form(default=None),
):
    # Parse file -> clean -> (later) extract skills -> (later) embed -> store raw+cleaned JSON for dev
    allowed = {"application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    if file.content_type not in allowed:
        # Fallback by extension if content-type is missing/incorrect
        filename = file.filename or ""
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        if ext not in {"pdf", "docx"}:
            raise HTTPException(status_code=415, detail=f"Unsupported content type: {file.content_type}")

    resume_id = str(uuid4())
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file uploaded")
    # Enforce max upload size (default 10 MB)
    max_bytes = get_max_upload_bytes()
    if len(data) > max_bytes:
        raise HTTPException(status_code=413, detail=f"File too large. Max allowed is {max_bytes // (1024*1024)} MB")

    raw_text = ""
    cleaned_text = ""
    meta: dict = {}

    is_pdf = file.content_type == "application/pdf" or (file.filename or "").lower().endswith(".pdf")
    is_docx = (
        file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or (file.filename or "").lower().endswith(".docx")
    )

    try:
        if is_pdf:
            raw_text, pages = extract_text_from_pdf(data)
            meta = {"type": "pdf", "pages": pages}
        elif is_docx:
            raw_text = extract_text_from_docx(data)
            meta = {"type": "docx"}
        else:
            raise HTTPException(status_code=415, detail="Unsupported file type. Only PDF and DOCX are allowed.")
    except Exception as e:
        # Tolerate parse failures: proceed with empty text, record parse error
        meta = {"type": "pdf" if is_pdf else ("docx" if is_docx else "unknown"), "parse_error": str(e)}
        raw_text = ""

    cleaned_text = clean_text_pipeline(raw_text)

    # Persist parsed output to JSON for local inspection
    try:
        save_parsed_json(resume_id, candidate_name, meta, raw_text, cleaned_text)
    except Exception:
        # Do not fail the request if local save fails; this is a dev convenience
        pass

    extracted_skills, _spans = extract_skills_with_spans(cleaned_text)

    # Embed resume text
    embedding = embed_text(cleaned_text)

    # persist to database
    insert_resume(
        resume_id=resume_id,
        candidate_name=candidate_name,
        raw_text=raw_text,
        cleaned_text=cleaned_text,
        skills=extracted_skills,
        embedding=embedding,
    )

    return UploadResumeResponse(
        resume_id=resume_id,
        candidate_name=candidate_name,
        skills=extracted_skills,
        embedding_dim=EMBEDDING_DIM,
        raw_text_len=len(raw_text or ""),
        cleaned_text_len=len(cleaned_text or ""),
        note=f"parsed: type={meta.get('type')}, raw_len={len(raw_text or '')}, cleaned_len={len(cleaned_text or '')}",
    )


@app.post("/upload_resumes", response_model=List[UploadResumeResponse], tags=["resumes"])  # type: ignore[arg-type]
async def upload_resumes(
    files: List[UploadFile] = File(..., description="One or more resume files: PDF or DOCX"),
    candidate_names: Optional[List[str]] = Form(default=None, description="Optional names aligned by index"),
):
    results: List[UploadResumeResponse] = []
    for idx, file in enumerate(files):
        try:
            # Validate type
            allowed = {"application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            filename = file.filename or ""
            ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
            is_pdf = file.content_type == "application/pdf" or ext == "pdf"
            is_docx = (
                file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == "docx"
            )
            if not (is_pdf or is_docx):
                raise HTTPException(status_code=415, detail=f"Unsupported content type: {file.content_type} ({filename})")

            # Read and size-check
            data = await file.read()
            if not data:
                raise HTTPException(status_code=400, detail=f"Empty file uploaded: {filename}")
            max_bytes = get_max_upload_bytes()
            if len(data) > max_bytes:
                raise HTTPException(status_code=413, detail=f"File too large: {filename}. Max {max_bytes // (1024*1024)} MB")

            # Parse
            meta: dict = {}
            if is_pdf:
                raw_text, pages = extract_text_from_pdf(data)
                meta = {"type": "pdf", "pages": pages}
            else:
                raw_text = extract_text_from_docx(data)
                meta = {"type": "docx"}

            cleaned_text = clean_text_pipeline(raw_text)
            cand_name = None
            if candidate_names and idx < len(candidate_names or []):
                cand_name = (candidate_names[idx] or "").strip() or None
            if not cand_name:
                # Default to filename without extension
                cand_name = filename.rsplit(".", 1)[0] if filename else None

            # Extract skills and embed
            skills = extract_skills(cleaned_text)
            embedding = embed_text(cleaned_text)

            resume_id = str(uuid4())

            # Save artifact
            try:
                save_parsed_json(resume_id, cand_name, meta, raw_text, cleaned_text)
            except Exception:
                pass

            # Persist
            insert_resume(
                resume_id=resume_id,
                candidate_name=cand_name,
                raw_text=raw_text,
                cleaned_text=cleaned_text,
                skills=skills,
                embedding=embedding,
            )

            results.append(
                UploadResumeResponse(
                    resume_id=resume_id,
                    candidate_name=cand_name,
                    skills=skills,
                    embedding_dim=EMBEDDING_DIM,
                    raw_text_len=len(raw_text or ""),
                    cleaned_text_len=len(cleaned_text or ""),
                    note=f"parsed: type={meta.get('type')}, raw_len={len(raw_text or '')}, cleaned_len={len(cleaned_text or '')}",
                )
            )
        except HTTPException as e:
            # Surface per-file error as a stub response with note
            results.append(
                UploadResumeResponse(
                    resume_id=str(uuid4()),
                    candidate_name=(candidate_names[idx] if candidate_names and idx < len(candidate_names or []) else None),
                    skills=[],
                    embedding_dim=EMBEDDING_DIM,
                    raw_text_len=0,
                    cleaned_text_len=0,
                    note=f"error: {e.status_code} {e.detail}",
                )
            )
        except Exception as e:
            results.append(
                UploadResumeResponse(
                    resume_id=str(uuid4()),
                    candidate_name=(candidate_names[idx] if candidate_names and idx < len(candidate_names or []) else None),
                    skills=[],
                    embedding_dim=EMBEDDING_DIM,
                    raw_text_len=0,
                    cleaned_text_len=0,
                    note=f"error: {e}",
                )
            )
    return results


@app.get("/resumes", response_model=List[ResumeListItem], tags=["resumes"])  # type: ignore[arg-type]
async def list_resumes(limit: int = 50, offset: int = 0):
    from sqlalchemy import text as _sql_text
    l = max(1, min(int(limit), 200))
    o = max(0, int(offset))
    with get_session() as s:
        rows = s.execute(
            _sql_text(
                """
                SELECT id, candidate_name, skills, created_at, updated_at
                FROM resumes
                ORDER BY created_at DESC
                LIMIT :limit OFFSET :offset
                """
            ),
            {"limit": l, "offset": o},
        ).all()
        out: List[ResumeListItem] = []
        for rid, cname, skills, created_at, updated_at in rows:
            out.append(
                ResumeListItem(
                    resume_id=str(rid),
                    candidate_name=cname,
                    skills=skills or [],
                    created_at=str(created_at),
                    updated_at=str(updated_at),
                )
            )
        return out


@app.delete("/resume/{resume_id}", tags=["resumes"])  # type: ignore[arg-type]
async def delete_resume(resume_id: str):
    from sqlalchemy import text as _sql_text
    with get_session() as s:
        res = s.execute(_sql_text("DELETE FROM resumes WHERE id=:id"), {"id": resume_id})
        s.commit()
        if res.rowcount and res.rowcount > 0:
            return {"deleted": True, "resume_id": resume_id}
        raise HTTPException(status_code=404, detail="resume not found")


@app.get("/resumes/{resume_id}", response_model=ResumeOut, tags=["resumes"])  # type: ignore[arg-type]
async def get_resume_text(resume_id: str):
    from sqlalchemy import text as _sql_text
    with get_session() as s:
        row = s.execute(
            _sql_text("SELECT id, candidate_name, cleaned_text, skills FROM resumes WHERE id=:id"), {"id": resume_id}
        ).first()
        if not row:
            raise HTTPException(status_code=404, detail="resume not found")
        rid, cname, cleaned, skills = row
        return ResumeOut(resume_id=str(rid), candidate_name=cname, cleaned_text=cleaned or "", skills=skills or [])


@app.get("/jobs/{job_id}", response_model=JobOut, tags=["jobs"])  # type: ignore[arg-type]
async def get_job_info(job_id: str):
    from sqlalchemy import text as _sql_text
    with get_session() as s:
        row = s.execute(
            _sql_text("SELECT id, title, description_cleaned, required_skills FROM jobs WHERE id=:id"),
            {"id": job_id},
        ).first()
        if not row:
            raise HTTPException(status_code=404, detail="job not found")
        jid, title, desc, req = row
        return JobOut(job_id=str(jid), title=title, description=desc or "", required_skills=req or [])


@app.get("/jobs", response_model=List[JobListItem], tags=["jobs"])  # type: ignore[arg-type]
async def list_jobs(limit: int = 50, offset: int = 0):
    from sqlalchemy import text as _sql_text
    l = max(1, min(int(limit), 200))
    o = max(0, int(offset))
    with get_session() as s:
        rows = s.execute(
            _sql_text(
                """
                SELECT id, title, required_skills, created_at, updated_at
                FROM jobs
                ORDER BY created_at DESC
                LIMIT :limit OFFSET :offset
                """
            ),
            {"limit": l, "offset": o},
        ).all()
        out: List[JobListItem] = []
        for jid, title, req, created_at, updated_at in rows:
            out.append(
                JobListItem(
                    job_id=str(jid),
                    title=title,
                    created_at=str(created_at),
                    updated_at=str(updated_at),
                    required_skills_count=len(req or []),
                )
            )
        return out


@app.delete("/job/{job_id}", tags=["jobs"])  # type: ignore[arg-type]
async def delete_job(job_id: str):
    from sqlalchemy import text as _sql_text
    with get_session() as s:
        res = s.execute(_sql_text("DELETE FROM jobs WHERE id=:id"), {"id": job_id})
        s.commit()
        if res.rowcount and res.rowcount > 0:
            return {"deleted": True, "job_id": job_id}
        raise HTTPException(status_code=404, detail="job not found")


@app.post("/upload_job", response_model=UploadJobResponse, tags=["jobs"])  # type: ignore[arg-type]
async def upload_job(payload: JobIn):
    # Clean and store job; extract required skills if not provided
    job_id = str(uuid4())
    cleaned_desc = clean_text_pipeline(payload.description)
    required = payload.required_skills or extract_skills(cleaned_desc)
    required = normalize_skill_list(required)
    job_embedding = embed_text(cleaned_desc)

    # persist to database
    insert_job(
        job_id=job_id,
        title=payload.title,
        description_cleaned=cleaned_desc,
        required_skills=required,
        embedding=job_embedding,
    )

    return UploadJobResponse(
        job_id=job_id,
        title=payload.title,
        required_skills=required,
        embedding_dim=EMBEDDING_DIM,
        note="stored in database; sentence-transformers enabled",
    )


@app.post("/upload_job_form", response_model=UploadJobResponse, tags=["jobs"])  # type: ignore[arg-type]
async def upload_job_form(
    title: str = Form(...),
    description: str = Form(...),
    required_skills: Optional[str] = Form(default=None, description="Comma or newline separated list"),
):
    job_id = str(uuid4())
    cleaned_desc = clean_text_pipeline(description)
    req_list = parse_skill_str_list(required_skills)
    required = normalize_skill_list(req_list) if req_list else extract_skills(cleaned_desc)
    job_embedding = embed_text(cleaned_desc)

    # persist to database
    insert_job(
        job_id=job_id,
        title=title,
        description_cleaned=cleaned_desc,
        required_skills=required,
        embedding=job_embedding,
    )

    return UploadJobResponse(
        job_id=job_id,
        title=title,
        required_skills=required,
        embedding_dim=EMBEDDING_DIM,
        note="stored in database; sentence-transformers enabled (form)",
    )




@app.get("/match", response_model=MatchResponse, tags=["matching"])  # type: ignore[arg-type]
async def match(job_id: str, k: int = 20):
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail=f"job_id '{job_id}' not found")

    req_skills = set(job.get("required_skills", []) or [])
    job_emb = job.get("embedding")
    # Avoid truthiness check on array-like objects (e.g., numpy arrays)
    if job_emb is None or (hasattr(job_emb, "__len__") and len(job_emb) == 0):
        return MatchResponse(job_id=job_id, k=k, results=[], note="job has no embedding")

    # Vector search in Postgres (pgvector cosine distance)
    topk = search_resumes_by_embedding(job_embedding=job_emb, k=k)

    results: list[MatchResult] = []
    # Fetch cleaned_text for matched span explanations
    from sqlalchemy import text as _sql_text
    job_desc = job.get("description_cleaned") or ""
    with get_session() as _s:
        for resume_id, candidate_name, res_skills_list, cosine in topk:
            res_skills = set(res_skills_list or [])
            overlap = jaccard(req_skills, res_skills) if req_skills else 0.0
            matched = sorted(req_skills & res_skills)
            missing = sorted(req_skills - res_skills)
            composite = 0.7 * float(cosine) + 0.3 * float(overlap)

            # Load resume text and compute spans, filter to matched skills
            cleaned_text_row = _s.execute(_sql_text("SELECT cleaned_text FROM resumes WHERE id=:id"), {"id": resume_id}).scalar()
            spans_filtered: list[dict] = []
            context_terms: List[str] = []
            context_job_spans: List[dict] = []
            context_resume_spans: List[dict] = []
            if isinstance(cleaned_text_row, str) and cleaned_text_row:
                _skills_all, spans_all = extract_skills_with_spans(cleaned_text_row)
                matched_set = set(matched)
                spans_filtered = [sp for sp in spans_all if sp.get("skill") in matched_set]
                # Compute context overlaps excluding skills
                context_terms = extract_context_terms(job_desc, cleaned_text_row, exclude_terms=matched, max_terms=20)
                context_job_spans = find_spans_for_terms(job_desc, context_terms)
                context_resume_spans = find_spans_for_terms(cleaned_text_row, context_terms)

            results.append(
                MatchResult(
                    resume_id=resume_id,
                    candidate_name=candidate_name,
                    cosine=float(cosine),
                    skills_overlap=float(overlap),
                    composite_score=float(composite),
                    matched_skills=matched,
                    missing_skills=missing,
                    matched_spans=spans_filtered,
                    context_terms=context_terms,
                    context_job_spans=context_job_spans,
                    context_resume_spans=context_resume_spans,
                )
            )

    # Sort and take top-k (already top-k by vector distance; still sort by composite)
    results.sort(key=lambda r: r.composite_score, reverse=True)
    results = results[: max(1, k)]

    return MatchResponse(job_id=job_id, k=k, results=results, note="pgvector cosine + skills jaccard")
